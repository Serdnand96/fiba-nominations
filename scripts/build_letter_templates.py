"""Build the placeholder (docxtpl) letter templates from the original .docx files.

The original templates are positional skeletons: the builders in
document_generator.py write into paras[2], paras[4], … so the file's paragraph
layout is load-bearing. This script converts them into real templates whose
body is {{ placeholders }}, keeping the letterhead, logos and signature block
of the source file untouched.

Run from the repo root:  python scripts/build_letter_templates.py
The generated *_TPL.docx files are committed. Re-run when a letterhead changes.
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

TEMPLATES = Path("templates")
DARK = RGBColor(0x2A, 0x2A, 0x2A)

RIGHT = WD_ALIGN_PARAGRAPH.RIGHT
CENTER = WD_ALIGN_PARAGRAPH.CENTER
JUSTIFY = WD_ALIGN_PARAGRAPH.JUSTIFY

# Body text shared by the nomination letters. Only the font, the paragraph the
# body starts at and the fee style differ between WCQ and GENERIC.
#   {{ x }}   plain value        {{r x }}  RichText (carries colour/bold)
#   {%p %}    paragraph-level Jinja tag, removed on render
def nomination_body(fee_style=None, blank_after_date=1):
    # Blank-line counts mirror the positional builders exactly; they are the
    # part a text-only diff cannot see, so they are spelled out here:
    #   date -> Dear      : GENERIC 1, WCQ 0 (its date sits at paras[1], Dear at [2])
    #   payment -> fees   : 2  (old: fee_idx = payment_idx + 3)
    #   last fee -> closing: 2 (old: closing_idx = fee_idx + len(fees) + 2)
    return [
        # (text, align, size_pt, style)
        ("{{r letter_date }}", RIGHT, 10, None),
        *[("", None, None, None)] * blank_after_date,
        ("{{r greeting }}", None, None, None),
        ("", None, None, None),
        ("We would like to inform that you have been nominated for the following "
         "games of the {{ competition }}.", None, None, None),
        ("", None, None, None),
        ("{%p for game in game_dates %}", None, None, None),
        ("{{r game }}", CENTER, 10, None),
        ("{%p endfor %}", None, None, None),
        ("{{r host }}", CENTER, 10, None),
        ("", None, None, None),
        ("{{r confirmation_paragraph }}", JUSTIFY, None, None),
        ("", None, None, None),
        ("As soon as we receive your confirmation, we will make arrangements for "
         "international flights to the host country and provide you with relevant "
         "information in order for you to prepare the game and establish contact "
         "with the Game Director of the Host National Federation.", None, None, None),
        ("", None, None, None),
        ("Below list the details of payment you will receive as {{ role }} "
         "assigned to the competition listed above:", None, None, None),
        ("", None, None, None),
        ("", None, None, None),
        ("{%p for fee in payment_lines %}", None, None, None),
        ("{{r fee }}", None, 10, fee_style),
        ("{%p endfor %}", None, None, None),
        ("", None, None, None),
        ("", None, None, None),
        ("We wish you the best in your preparation and accomplishment of your "
         "assignment.", None, None, None),
        # Gap before the signature. The positional builders sized this
        # dynamically (`keep_empty = max(0, 30 - closing_idx - 2)`), which came
        # out to 5-7 blank lines on a typical 1-3 game letter. A declarative
        # template can't do that arithmetic, so it uses a fixed 5.
        ("", None, None, None),
        ("", None, None, None),
        ("", None, None, None),
        ("", None, None, None),
        ("", None, None, None),
    ]


# BCLA is a confirmation letter, not a nomination: different wording, fees in
# ink instead of red, and two paragraphs whose text depends on the F4/RS
# variant (resolved in _bcla_context, so the .docx stays a flat sequence).
# Labelled lines use {%p if %} so the wording stays editable in Word.
BCLA_BODY = [
    ("{{r heading }}", None, None, None),
    ("", None, None, None),
    ("{{r greeting }}", None, None, None),
    ("", None, None, None),
    ("By way of this letter, we confirm your acceptance for your assignment as "
     "{{ role }} for the {{ competition }} {{ year }}.",
     JUSTIFY, 10, None),
    ("", None, None, None),
    ("Game Information", JUSTIFY, 10, None),          # bold applied below
    ("{%p if location %}", None, None, None),
    ("Location: {{ location }}.", JUSTIFY, 10, None),
    ("{%p endif %}", None, None, None),
    ("{%p if venue %}", None, None, None),
    ("Venue: {{ venue }}", JUSTIFY, 10, None),
    ("{%p endif %}", None, None, None),
    ("", None, None, None),
    ("{%p if arrival_date %}", None, None, None),
    ("Arrival Date: {{ arrival_date }}", JUSTIFY, 10, None),
    ("{%p endif %}", None, None, None),
    ("{%p for game in game_dates %}", None, None, None),
    ("{{ game }}", JUSTIFY, 10, None),
    ("{%p endfor %}", None, None, None),
    ("{%p if departure_date %}", None, None, None),
    ("Departure Date: {{ departure_date }}", JUSTIFY, 10, None),
    ("{%p endif %}", None, None, None),
    ("", None, None, None),
    ("Financial Details", JUSTIFY, 10, None),         # bold applied below
    ("{{ payment_intro }}", JUSTIFY, 10, None),
    ("{%p for fee in payment_lines %}", None, None, None),
    ("{{r fee }}", JUSTIFY, None, None),
    ("{%p endfor %}", None, None, None),
    ("", None, None, None),
    ("Additionally, breakfast, lunch and dinner will be provided by the club at "
     "your hotel as per the dates of your assigned games.", JUSTIFY, 10, None),
    ("", None, None, None),
    ("Payment for this assignment will be made within 21-days of the window "
     "conclusion. ", JUSTIFY, 10, None),
    ("", None, None, None),
    ("{{ banking_paragraph }}", JUSTIFY, 10, None),
    ("", None, None, None),
    ("If you have any questions, please do not hesitate to contact.",
     JUSTIFY, 10, None),
    ("", None, None, None),
]

# Headings the original builder emitted in bold.
BCLA_BOLD = {"Game Information", "Financial Details"}


SPECS = {
    "GENERIC": {
        "src": "GENERIC_TEMPLATE.docx",
        "dst": "GENERIC_TEMPLATE_TPL.docx",
        "font": "Univers",
        # Body starts at the very first paragraph: the letterhead logo lives in
        # the section header, not in the body.
        "body_start": 0,
        "body": nomination_body(),
    },
    "WCQ": {
        "src": "WCQ_TEMPLATE.docx",
        "dst": "WCQ_TEMPLATE_TPL.docx",
        "font": "IBM Plex Sans",
        # [0] is the title + competition logo — preserved as-is.
        "body_start": 1,
        "body": nomination_body(fee_style="List Paragraph", blank_after_date=0),
    },
    "BCLA": {
        "src": "BCLA_TEMPLATE.docx",
        "dst": "BCLA_TEMPLATE_TPL.docx",
        "font": "Univers",
        # BCLA's source has only 15 paragraphs — [0][1] logos, [2] the date,
        # [4] the single content paragraph, then "Respectfully," and the
        # signature. There is nothing to overwrite, so the body is inserted
        # after the anchor and the tail is left untouched.
        "anchor": 4,
        "date_para": 2,
        "date_tag": "{{r letter_date }}",
        "body": BCLA_BODY,
        "bold": BCLA_BOLD,
    },
    "LSB": {
        # No source file: the letter was built entirely in code.
        "scratch": True,
        "dst": "LSB_TEMPLATE_TPL.docx",
        "font": "IBM Plex Sans",
    },
}


def build_lsb(name, spec):
    """Create the LSB template from nothing.

    LSB has no base .docx — _build_confirmation_from_scratch builds the whole
    letter in code. To guarantee the template's formatting matches byte for
    byte, this reuses that builder's own helpers and feeds them Jinja tags
    instead of values.
    """
    import sys
    sys.path.insert(0, ".")
    from api._lib.services.document_generator import (
        _apply_base_style, _add_heading, _add_body, _add_body_text,
        _add_centered_red, _add_fee_line, _add_empty, COLOR_DARK,
    )

    doc = Document()
    _apply_base_style(doc)

    def tag(text):
        """A Jinja control tag: its formatting is irrelevant, docxtpl drops it."""
        _add_body_text(doc, text)

    _add_heading(doc, "{{ heading }}")
    _add_empty(doc)
    _add_body(doc, [("{{r greeting }}", COLOR_DARK)])
    _add_empty(doc)
    _add_body_text(doc, "This letter confirms your assignment as {{ role }} "
                        "for the {{ competition }}.")
    _add_empty(doc)

    # Detail bullets — each guarded so the line disappears when the value is
    # absent, exactly like the `if data.get(...)` chain in the builder.
    for field, label in (("location", "Location"), ("venue", "Venue"),
                         ("arrival_date", "Arrival Date"),
                         ("departure_date", "Departure Date")):
        tag("{%p if " + field + " %}")
        _add_body_text(doc, f"  •  {label}: {{{{ {field} }}}}")
        tag("{%p endif %}")
    _add_empty(doc)

    tag("{%p for game in game_dates %}")
    _add_centered_red(doc, "{{r game }}")
    tag("{%p endfor %}")
    _add_empty(doc)

    _add_body_text(doc, "Below list the details of payment you will receive as "
                        "{{ role }} assigned to the competition listed above:")
    _add_empty(doc)
    tag("{%p for fee in payment_lines %}")
    _add_fee_line(doc, "{{r fee }}")
    tag("{%p endfor %}")
    _add_empty(doc)

    _add_body_text(doc, "Thank you for your commitment and professionalism.")
    _add_empty(doc)
    _add_empty(doc)
    _add_body_text(doc, "{{ signature }}")

    dst = TEMPLATES / spec["dst"]
    doc.save(str(dst))
    print(f"{name}: built from scratch -> {dst.name} "
          f"({len(Document(str(dst)).paragraphs)} paras)")


def write_para(para, text, align, size, style_name, font, doc, bold=False):
    """Replace a paragraph's runs with a single formatted run."""
    for run in list(para.runs):
        run._element.getparent().remove(run._element)
    if align is not None:
        para.alignment = align
    if style_name:
        try:
            para.style = doc.styles[style_name]
        except KeyError:
            pass
    if text:
        run = para.add_run(text)
        run.font.name = font
        run.font.color.rgb = DARK
        run.bold = bold
        if size:
            run.font.size = Pt(size)


def build_insert(name, spec):
    """Build a template whose body has to be inserted, not overwritten.

    Mirrors what _build_bcla_letter does at render time: the first line goes
    into the anchor paragraph and the rest are inserted after it, so the
    letterhead above and the signature below stay exactly where they are.
    """
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph

    src = TEMPLATES / spec["src"]
    doc = Document(str(src))
    font = spec["font"]
    bold_set = spec.get("bold", set())

    for style_name in ("Normal", "Body Text", "Cuerpo"):
        try:
            doc.styles[style_name].font.name = font
        except KeyError:
            pass

    paras = doc.paragraphs

    if "date_para" in spec:
        write_para(paras[spec["date_para"]], spec["date_tag"], RIGHT, 10, None, font, doc)

    body = spec["body"]
    anchor = paras[spec["anchor"]]
    first_text, first_align, first_size, first_style = body[0]
    write_para(anchor, first_text, first_align, first_size, first_style, font, doc,
               bold=first_text in bold_set)

    insert_after = anchor._element
    for text, align, size, style_name in body[1:]:
        new_p = doc.element.makeelement(qn("w:p"), {})
        insert_after.addnext(new_p)
        insert_after = new_p
        write_para(Paragraph(new_p, doc), text, align, size, style_name, font, doc,
                   bold=text in bold_set)

    dst = TEMPLATES / spec["dst"]
    doc.save(str(dst))
    out = Document(str(dst))
    imgs = [i for i, p in enumerate(out.paragraphs) if "graphic" in p._p.xml]
    print(f"{name}: {src.name} ({len(paras)} paras) -> {dst.name} "
          f"({len(out.paragraphs)} paras, inserted), images at {imgs}")


def build(name, spec):
    if spec.get("scratch"):
        return build_lsb(name, spec)
    if "anchor" in spec:
        return build_insert(name, spec)

    src = TEMPLATES / spec["src"]
    doc = Document(str(src))
    font = spec["font"]

    # Match what the positional builders did: set the default font, and only
    # the font. Pinning the default *colour* here would also repaint the
    # signature block and the footer, which the originals leave pure black.
    # Mixed-ink paragraphs are handled as whole RichText values instead.
    for style_name in ("Normal", "Body Text", "Heading 1"):
        try:
            doc.styles[style_name].font.name = font
        except KeyError:
            pass

    paras = doc.paragraphs

    # Find the signature block by locating its scanned image rather than using
    # a fixed offset from the end — an offset silently eats the image when a
    # letterhead gains or loses a trailing paragraph.
    images = [i for i, p in enumerate(paras) if "graphic" in p._p.xml]
    tail_images = [i for i in images if i > spec["body_start"]]
    if not tail_images:
        raise SystemExit(f"{name}: no signature image found after the body start")
    sig_start = tail_images[-1]

    body = paras[spec["body_start"]:sig_start]
    content = spec["body"]
    if len(content) > len(body):
        raise SystemExit(
            f"{name}: template too short — need {len(content)} body paragraphs, "
            f"have {len(body)}")

    for para, (text, align, size, style_name) in zip(body, content):
        for run in list(para.runs):
            run._element.getparent().remove(run._element)
        if align is not None:
            para.alignment = align
        if style_name:
            try:
                para.style = doc.styles[style_name]
            except KeyError:
                pass
        if text:
            run = para.add_run(text)
            run.font.name = font
            run.font.color.rgb = DARK
            if size:
                run.font.size = Pt(size)

    # Drop leftover empties so the signature isn't pushed to a second page —
    # the Jinja loops expand at render time instead.
    for para in body[len(content):]:
        para._element.getparent().remove(para._element)

    dst = TEMPLATES / spec["dst"]
    doc.save(str(dst))

    out = Document(str(dst))
    sig = [i for i, p in enumerate(out.paragraphs) if "graphic" in p._p.xml]
    print(f"{name}: {src.name} ({len(paras)} paras) -> {dst.name} "
          f"({len(out.paragraphs)} paras), images at {sig}")


if __name__ == "__main__":
    for name, spec in SPECS.items():
        build(name, spec)
