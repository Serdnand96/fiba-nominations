"""Build templates/GENERIC_TEMPLATE_TPL.docx from GENERIC_TEMPLATE.docx.

Keeps the letterhead (header logo, footer address) and the signature block,
and replaces the 39 positional body paragraphs with docxtpl/Jinja placeholders.

Run once; the produced .docx is committed. Re-run if the letterhead changes.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = Path("templates/GENERIC_TEMPLATE.docx")
DST = Path("templates/GENERIC_TEMPLATE_TPL.docx")

FONT = "Univers"
DARK = RGBColor(0x2A, 0x2A, 0x2A)

# (text, align, size_pt, bold)  — one entry per body paragraph.
# {{ }} placeholders are plain text; {{r }} are RichText (carry their own
# colour/bold), and {%p %} lines are paragraph-level Jinja control tags that
# docxtpl removes on render.
BODY = [
    ("{{ letter_date }}", WD_ALIGN_PARAGRAPH.RIGHT, 10, False),
    ("", None, None, False),
    ("Dear {{r nominee }},", None, None, False),
    ("", None, None, False),
    ("We would like to inform that you have been nominated for the following "
     "games of the {{ competition_name }}.", None, None, False),
    ("", None, None, False),
    ("{%p for game in game_dates %}", None, None, False),
    ("{{r game }}", WD_ALIGN_PARAGRAPH.CENTER, 10, False),
    ("{%p endfor %}", None, None, False),
    ("{{r host_line }}", WD_ALIGN_PARAGRAPH.CENTER, 10, False),
    ("", None, None, False),
    ("As per the FIBA Internal Regulations Book 3, please confirm to us your "
     "availability to fulfil your assignment as {{ role_label }} by {{r deadline }}. "
     "Confirmation shall be sent to {{ confirm_email }}",
     WD_ALIGN_PARAGRAPH.JUSTIFY, None, False),
    ("", None, None, False),
    ("As soon as we receive your confirmation, we will make arrangements for "
     "international flights to the host country and provide you with relevant "
     "information in order for you to prepare the game and establish contact "
     "with the Game Director of the Host National Federation.", None, None, False),
    ("", None, None, False),
    ("Below list the details of payment you will receive as {{ role_label }} "
     "assigned to the competition listed above:", None, None, False),
    ("", None, None, False),
    ("{%p for fee in fee_lines %}", None, None, False),
    ("{{r fee }}", None, 10, False),
    ("{%p endfor %}", None, None, False),
    ("", None, None, False),
    ("We wish you the best in your preparation and accomplishment of your "
     "assignment.", None, None, False),
    # Gap before the signature. The positional builder computed this
    # dynamically (`keep_empty = max(0, 30 - closing_idx - 2)`) to pin the
    # signature near paragraph 30, which worked out to 5-7 blank lines for a
    # typical 1-3 game letter. A declarative template can't do that
    # arithmetic, so it uses a fixed 5 — matching the common 3-game case.
    ("", None, None, False),
    ("", None, None, False),
    ("", None, None, False),
    ("", None, None, False),
    ("", None, None, False),
]


def main():
    doc = Document(str(SRC))

    # docxtpl splits the run that holds a tag, and the text trailing a {{r }}
    # insert comes back without explicit colour. Pin the document default so
    # inherited text renders in the same ink as the explicit runs.
    for style_name in ("Normal", "Body Text", "Heading 1"):
        try:
            style = doc.styles[style_name]
            style.font.name = FONT
            style.font.color.rgb = DARK
        except KeyError:
            pass

    paras = doc.paragraphs

    # The signature block is the tail of the document: the scanned signature
    # image followed by the name/title lines. Find it by locating the image
    # rather than by a fixed offset — an offset silently eats the image if the
    # letterhead ever gains or loses a trailing paragraph.
    img_idx = [i for i, p in enumerate(paras) if "graphic" in p._p.xml]
    if not img_idx:
        raise SystemExit("no signature image found in the source template")
    sig_start = img_idx[-1]
    print(f"source paragraphs: {len(paras)}, signature image at {sig_start}")

    body = paras[:sig_start]
    if len(BODY) > len(body):
        raise SystemExit(f"template too short: need {len(BODY)}, have {len(body)}")

    for para, (text, align, size, bold) in zip(body, BODY):
        for run in list(para.runs):
            run._element.getparent().remove(run._element)
        if align is not None:
            para.alignment = align
        if text:
            run = para.add_run(text)
            run.font.name = FONT
            run.font.color.rgb = DARK
            run.bold = bold
            if size:
                run.font.size = Pt(size)

    # Drop the leftover empty paragraphs so the signature isn't pushed onto a
    # second page — the Jinja loops expand at render time instead.
    for para in body[len(BODY):]:
        para._element.getparent().remove(para._element)

    doc.save(str(DST))
    out = Document(str(DST))
    print(f"wrote {DST} — {len(out.paragraphs)} paragraphs")
    for i, p in enumerate(out.paragraphs):
        if p.text.strip():
            print(f"  [{i}] {p.text[:72]}")


if __name__ == "__main__":
    main()
