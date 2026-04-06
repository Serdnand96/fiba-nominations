import os
import re
import copy
import tempfile
import httpx
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUTPUT_DIR = Path(tempfile.gettempdir()) / "fiba_generated"
TEMPLATES_DIR = Path(__file__).resolve().parent.parent.parent.parent / "templates"

# FIBA brand colors
COLOR_DARK = RGBColor(0x2A, 0x2A, 0x2A)
COLOR_RED = RGBColor(0xED, 0x00, 0x00)

# FIBA brand fonts per template
FONT_WCQ = "IBM Plex Sans"
FONT_GENERIC = "Univers"
FONT_NAME = FONT_WCQ  # default

CONFIRMATION_EMAIL = {
    "VGO": "vgo.americas@fiba.basketball",
    "TD": "competitions-americas@fiba.basketball",
}

SIGNATORIES = {
    "WCQ": ("Carlos Alves", "Executive Director", "FIBA Americas"),
    "GENERIC": ("Carlos Alves", "Executive Director", "FIBA Americas"),
    "BCLA_F4": ("Gino Rullo", "Head of Operations", "Basketball Champions League Americas"),
    "BCLA_RS": ("Gino Rullo", "Head of Operations", "Basketball Champions League Americas"),
    "LSB": ("Gino Rullo", "Head of Operations", "Club Competitions – FIBA Americas"),
}


def generate_nomination(nomination_data: dict) -> tuple[str, str | None, str | None]:
    """Generate a nomination/confirmation .docx letter, convert to PDF, upload.
    Returns (local_path, storage_url, conversion_error).
    """
    template_key = nomination_data["template_key"]

    if template_key == "WCQ":
        doc = _build_wcq_letter(nomination_data)
    elif template_key == "GENERIC":
        doc = _build_generic_letter(nomination_data)
    elif template_key == "BCLA_F4":
        doc = _build_bcla_letter(nomination_data, variant="F4")
    elif template_key == "BCLA_RS":
        doc = _build_bcla_letter(nomination_data, variant="RS")
    elif template_key == "LSB":
        doc = _build_confirmation_from_scratch(nomination_data)
    else:
        raise ValueError(f"Unknown template_key: {template_key}")

    # Build output filename: "Nombre Apellido Competencia Nomination"
    name_clean = re.sub(r"[^\w\s-]", "", nomination_data["nominee_name"]).strip()
    comp_clean = re.sub(r"[^\w\s-]", "", nomination_data["competition_name"]).strip()
    base_name = f"{name_clean} {comp_clean} Nomination"

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    docx_path = OUTPUT_DIR / f"{base_name}.docx"
    doc.save(str(docx_path))

    # Convert to PDF
    pdf_path, conversion_error = _convert_to_pdf(str(docx_path))
    final_path = pdf_path if pdf_path else str(docx_path)

    # Upload to Supabase Storage
    storage_url = _upload_to_storage(final_path, base_name)
    return final_path, storage_url, conversion_error


# ─── DATE FORMATTING ─────────────────────────────────────────────────────────

def _fmt_date(date_str: str) -> str:
    """Convert ISO date (2026-04-17) to readable format (17 April 2026)."""
    if not date_str:
        return ""
    try:
        dt = datetime.strptime(date_str, "%Y-%m-%d")
        return dt.strftime("%-d %B %Y")
    except (ValueError, TypeError):
        try:
            dt = datetime.strptime(date_str, "%Y-%m-%d")
            return dt.strftime("%d %B %Y").lstrip("0")
        except Exception:
            return date_str


def _fmt_deadline(date_str: str) -> str:
    """Format deadline: January 18th, 2026."""
    if not date_str:
        return ""
    try:
        dt = datetime.strptime(date_str, "%Y-%m-%d")
        day = dt.day
        if 11 <= day <= 13:
            suffix = "th"
        else:
            suffix = {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
        return dt.strftime(f"%B {day}{suffix}, %Y")
    except Exception:
        return date_str


# ─── WCQ / GENERIC LETTER ────────────────────────────────────────────────────

def _build_wcq_letter(data: dict) -> Document:
    template_path = TEMPLATES_DIR / "WCQ_TEMPLATE.docx"
    if not template_path.exists():
        return _build_wcq_from_scratch(data)

    doc = Document(str(template_path))

    # Set IBM Plex Sans as the default font for the document
    for style_name in ["Normal", "Body Text", "Heading 1"]:
        try:
            doc.styles[style_name].font.name = FONT_NAME
        except Exception:
            pass

    paras = doc.paragraphs

    nominee = data.get("nominee_name", "")
    comp_name = data.get("competition_name", "")
    role = data.get("role", "VGO")
    role_label = "Video Graphic Operator" if role == "VGO" else "Technical Delegate"
    game_dates = data.get("game_dates") or []
    deadline = _fmt_deadline(data.get("confirmation_deadline", ""))
    letter_date = _fmt_date(data.get("letter_date", ""))
    fee = data.get("window_fee")
    incidentals = data.get("incidentals")
    total = data.get("total")

    # Clear paragraphs 1-44 (keep [0] title+logo, [45] signature image, [46] signature text)
    for i in range(1, len(paras) - 2):
        _clear_para(paras[i])

    # [1] — Letter date (right-aligned, red)
    if letter_date:
        _set_para_text(paras[1], letter_date, COLOR_RED, size=Pt(10),
                      align=WD_ALIGN_PARAGRAPH.RIGHT)

    # [2] — "Dear [Name],"
    _set_para_mixed(paras[2], [
        ("Dear ", COLOR_DARK, False),
        (nominee, COLOR_RED, False),
        (",", COLOR_DARK, False),
    ])

    # [4] — Body intro
    _set_para_text(paras[4],
        f"We would like to inform that you have been nominated for the "
        f"following games of the {comp_name}.",
        COLOR_DARK)

    # [6+] — Game dates (centered, bold, red)
    for i, gd in enumerate(game_dates):
        idx = 6 + i
        if idx < len(paras) - 3:
            label = gd.get("label", "")
            date_val = _fmt_date(gd.get("date", ""))
            text = f"{label}: {date_val}" if label else date_val
            _set_para_text(paras[idx], text, COLOR_RED, bold=True, size=Pt(10),
                          align=WD_ALIGN_PARAGRAPH.CENTER)

    # Confirmation paragraph — 2 lines after last game date
    confirm_email = CONFIRMATION_EMAIL.get(role, CONFIRMATION_EMAIL["VGO"])
    confirm_idx = 6 + max(len(game_dates), 1) + 2
    if confirm_idx < len(paras) - 3:
        _set_para_mixed(paras[confirm_idx], [
            (f"As per the FIBA Internal Regulations Book 3, please confirm to us "
             f"your availability to fulfil your assignment as {role_label} by ",
             COLOR_DARK, False),
            (f"{deadline}", COLOR_RED, False),
            (".", COLOR_DARK, False),
            (" Confirmation shall be sent to ", COLOR_DARK, False),
            (confirm_email, COLOR_DARK, False),
        ], align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # Travel paragraph
    travel_idx = confirm_idx + 2
    if travel_idx < len(paras) - 3:
        _set_para_text(paras[travel_idx],
            "As soon as we receive your confirmation, we will make arrangements "
            "for international flights to the host country and provide you with "
            "relevant information in order for you to prepare the game and "
            "establish contact with the Game Director of the Host National Federation.",
            COLOR_DARK)

    # Payment intro
    payment_idx = travel_idx + 2
    if payment_idx < len(paras) - 3:
        _set_para_text(paras[payment_idx],
            f"Below list the details of payment you will receive as {role_label} "
            f"assigned to the competition listed above:",
            COLOR_DARK)

    # Fee items — 2 blank lines after payment intro, then the fees
    fee_idx = payment_idx + 3
    fee_items = [
        (f"Per Game Fee: {_fmt_money(fee)}", False),
        (f"Incidentals: {_fmt_money(incidentals)}", False),
        (f"Total: {_fmt_money(total)}", True),
    ]
    for i, (text, bold) in enumerate(fee_items):
        idx = fee_idx + i
        if idx < len(paras) - 3:
            _set_para_text(paras[idx], text, COLOR_RED, bold=bold, size=Pt(10))
            try:
                paras[idx].style = doc.styles["List Paragraph"]
            except Exception:
                pass

    # Closing — 2 blank lines after last fee
    closing_idx = fee_idx + len(fee_items) + 2
    if closing_idx < len(paras) - 3:
        _set_para_text(paras[closing_idx],
            "We wish you the best in your preparation and accomplishment of your assignment.",
            COLOR_DARK)

    # Remove excess empty paragraphs between closing and signature,
    # but keep enough so the signature sits at the bottom of the page.
    # Target: ~30 total paragraphs keeps signature near the bottom of page 1.
    keep_empty = max(0, 30 - closing_idx - 2)  # 2 = signature image + text
    remove_from = closing_idx + 1 + keep_empty
    if remove_from < len(paras) - 2:
        _remove_excess_paragraphs(doc, remove_from, len(paras) - 2)

    return doc


def _remove_excess_paragraphs(doc, start_idx, end_idx):
    """Remove empty paragraphs from start_idx to end_idx (exclusive) to compact the document."""
    body = doc.element.body
    paras = doc.paragraphs
    WP_DRAWING = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'
    W_DRAWING = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'
    to_remove = []
    for i in range(end_idx - 1, start_idx - 1, -1):
        p = paras[i]
        if not p.text.strip():
            # Check for any drawing elements (images) using direct XML search
            has_drawing = (
                len(p._element.findall(f'.//{W_DRAWING}')) > 0 or
                len(p._element.findall(f'.//{WP_DRAWING}anchor')) > 0 or
                len(p._element.findall(f'.//{WP_DRAWING}inline')) > 0
            )
            if not has_drawing:
                to_remove.append(p._element)
    for elem in to_remove:
        body.remove(elem)


# ─── GENERIC LETTER (Univers Condensed font, separate template) ──────────────

def _build_generic_letter(data: dict) -> Document:
    template_path = TEMPLATES_DIR / "GENERIC_TEMPLATE.docx"
    if not template_path.exists():
        return _build_wcq_from_scratch(data)

    doc = Document(str(template_path))
    font_name = FONT_GENERIC

    # Set default font for the document
    for style_name in ["Normal", "Body Text", "Heading 1"]:
        try:
            doc.styles[style_name].font.name = font_name
        except Exception:
            pass

    paras = doc.paragraphs

    nominee = data.get("nominee_name", "")
    comp_name = data.get("competition_name", "")
    role = data.get("role", "VGO")
    role_label = "Video Graphic Operator" if role == "VGO" else "Technical Delegate"
    game_dates = data.get("game_dates") or []
    deadline = _fmt_deadline(data.get("confirmation_deadline", ""))
    letter_date = _fmt_date(data.get("letter_date", ""))
    fee = data.get("window_fee")
    incidentals = data.get("incidentals")
    total = data.get("total")

    # Generic template structure: [0-38] content area, [39] signature image, [40-42] sig text
    # Clear content paragraphs (preserve signature at end)
    sig_start = len(paras) - 4  # last 4 paragraphs are signature area
    for i in range(0, sig_start):
        _clear_para_generic(paras[i])

    # [0] — Letter date (right-aligned)
    if letter_date:
        _set_para_text_font(paras[0], letter_date, COLOR_DARK, font_name, size=Pt(10),
                            align=WD_ALIGN_PARAGRAPH.RIGHT)

    # [2] — "Dear [Name],"
    _set_para_mixed_font(paras[2], [
        ("Dear ", COLOR_DARK, False),
        (nominee, COLOR_RED, False),
        (",", COLOR_DARK, False),
    ], font_name)

    # [4] — Body intro
    _set_para_text_font(paras[4],
        f"We would like to inform that you have been nominated for the "
        f"following games of the {comp_name}.",
        COLOR_DARK, font_name)

    # [6+] — Game dates (centered, bold, red)
    for i, gd in enumerate(game_dates):
        idx = 6 + i
        if idx < sig_start - 10:
            label = gd.get("label", "")
            date_val = _fmt_date(gd.get("date", ""))
            text = f"{label}: {date_val}" if label else date_val
            _set_para_text_font(paras[idx], text, COLOR_RED, font_name, bold=True, size=Pt(10),
                                align=WD_ALIGN_PARAGRAPH.CENTER)

    # Confirmation paragraph
    confirm_email = CONFIRMATION_EMAIL.get(role, CONFIRMATION_EMAIL["VGO"])
    confirm_idx = 6 + max(len(game_dates), 1) + 2
    if confirm_idx < sig_start - 8:
        _set_para_mixed_font(paras[confirm_idx], [
            (f"As per the FIBA Internal Regulations Book 3, please confirm to us "
             f"your availability to fulfil your assignment as {role_label} by ",
             COLOR_DARK, False),
            (f"{deadline}", COLOR_RED, False),
            (".", COLOR_DARK, False),
            (" Confirmation shall be sent to ", COLOR_DARK, False),
            (confirm_email, COLOR_DARK, False),
        ], font_name, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # Travel paragraph
    travel_idx = confirm_idx + 2
    if travel_idx < sig_start - 6:
        _set_para_text_font(paras[travel_idx],
            "As soon as we receive your confirmation, we will make arrangements "
            "for international flights to the host country and provide you with "
            "relevant information in order for you to prepare the game and "
            "establish contact with the Game Director of the Host National Federation.",
            COLOR_DARK, font_name)

    # Payment intro
    payment_idx = travel_idx + 2
    if payment_idx < sig_start - 5:
        _set_para_text_font(paras[payment_idx],
            f"Below list the details of payment you will receive as {role_label} "
            f"assigned to the competition listed above:",
            COLOR_DARK, font_name)

    # Fee items
    fee_idx = payment_idx + 3
    fee_items = [
        (f"Per Game Fee: {_fmt_money(fee)}", False),
        (f"Incidentals: {_fmt_money(incidentals)}", False),
        (f"Total: {_fmt_money(total)}", True),
    ]
    for i, (text, bold) in enumerate(fee_items):
        idx = fee_idx + i
        if idx < sig_start - 2:
            _set_para_text_font(paras[idx], text, COLOR_RED, font_name, bold=bold, size=Pt(10))

    # Closing
    closing_idx = fee_idx + len(fee_items) + 2
    if closing_idx < sig_start - 1:
        _set_para_text_font(paras[closing_idx],
            "We wish you the best in your preparation and accomplishment of your assignment.",
            COLOR_DARK, font_name)

    # Remove excess empty paragraphs between closing and signature
    keep_empty = max(0, 30 - closing_idx - 2)
    remove_from = closing_idx + 1 + keep_empty
    if remove_from < len(paras) - 4:
        _remove_excess_paragraphs(doc, remove_from, len(paras) - 4)

    return doc


# ─── BCLA LETTER (Univers font, BCLA template) ──────────────────────────────

def _build_bcla_letter(data: dict, variant: str = "F4") -> Document:
    template_path = TEMPLATES_DIR / "BCLA_TEMPLATE.docx"
    if not template_path.exists():
        return _build_confirmation_from_scratch(data)

    doc = Document(str(template_path))
    font_name = "Univers"

    for style_name in ["Normal", "Body Text"]:
        try:
            doc.styles[style_name].font.name = font_name
        except Exception:
            pass

    paras = doc.paragraphs

    nominee = data.get("nominee_name", "")
    comp_name = data.get("competition_name", "")
    comp_year = data.get("competition_year", "")
    role = data.get("role", "VGO")
    role_label = "Video Graphic Operator" if role == "VGO" else "Technical Delegate"
    game_dates = data.get("game_dates") or []
    location = data.get("location", "")
    venue = data.get("venue", "")
    arrival_date = data.get("arrival_date", "")
    departure_date = data.get("departure_date", "")
    letter_date = data.get("letter_date", "")
    fee = data.get("window_fee")
    incidentals = data.get("incidentals")
    total = data.get("total")

    # Format letter date like "Miami, March 27th, 2024"
    formatted_letter_date = ""
    if letter_date:
        formatted_letter_date = f"Miami, {_fmt_deadline(letter_date)}"

    # Clear the placeholder paragraph [4]
    _clear_para(paras[4])

    # Set date in paragraph [2] (right-aligned)
    if formatted_letter_date:
        _set_para_text_font(paras[2], formatted_letter_date, COLOR_DARK, font_name,
                            size=Pt(10), align=WD_ALIGN_PARAGRAPH.RIGHT)

    content_lines = []

    # Title
    content_lines.append({
        "text": f"BCL Americas {comp_year} – {role_label.upper()} NOMINATION",
        "bold": True, "color": COLOR_DARK, "size": Pt(11)
    })
    content_lines.append({"text": ""})

    # Dear
    content_lines.append({
        "mixed": [("Dear ", COLOR_DARK, False), (nominee, COLOR_RED, False), (",", COLOR_DARK, False)]
    })
    content_lines.append({"text": ""})

    # Confirmation body
    content_lines.append({
        "text": f"By way of this letter, we confirm your acceptance for your assignment as "
                f"{role_label} for the {comp_name} {comp_year}.",
        "color": COLOR_DARK, "align": WD_ALIGN_PARAGRAPH.JUSTIFY
    })
    content_lines.append({"text": ""})

    # Game Information
    content_lines.append({"text": "Game Information", "bold": True, "color": COLOR_DARK,
                          "align": WD_ALIGN_PARAGRAPH.JUSTIFY})
    if location:
        content_lines.append({"text": f"Location: {location}.", "color": COLOR_DARK,
                              "align": WD_ALIGN_PARAGRAPH.JUSTIFY})
    if venue:
        content_lines.append({"text": f"Venue: {venue}", "color": COLOR_DARK,
                              "align": WD_ALIGN_PARAGRAPH.JUSTIFY})
    content_lines.append({"text": ""})

    if arrival_date:
        content_lines.append({"text": f"Arrival Date: {_fmt_deadline(arrival_date)}",
                              "color": COLOR_DARK, "align": WD_ALIGN_PARAGRAPH.JUSTIFY})

    # F4 has game date rows (Semifinals, 3rd Place, Final); RS does not
    if variant == "F4":
        for gd in game_dates:
            label = gd.get("label", "")
            date_val = _fmt_deadline(gd.get("date", ""))
            text = f"{label}: {date_val}" if label else date_val
            content_lines.append({"text": text, "color": COLOR_DARK,
                                  "align": WD_ALIGN_PARAGRAPH.JUSTIFY})

    if departure_date:
        content_lines.append({"text": f"Departure Date: {_fmt_deadline(departure_date)}",
                              "color": COLOR_DARK, "align": WD_ALIGN_PARAGRAPH.JUSTIFY})

    content_lines.append({"text": ""})

    # Financial Details
    content_lines.append({"text": "Financial Details", "bold": True, "color": COLOR_DARK,
                          "align": WD_ALIGN_PARAGRAPH.JUSTIFY})

    if variant == "RS":
        content_lines.append({
            "text": f"Below lists the details of payment you will receive as a BCL Americas "
                    f"{role_label} assigned to the games listed above. The distribution of this "
                    f"payment is as follows:",
            "color": COLOR_DARK, "align": WD_ALIGN_PARAGRAPH.JUSTIFY
        })
    else:
        content_lines.append({
            "text": f"Below lists the details of payment you will receive as a BCL Americas "
                    f"{role_label} assigned to the games listed above:",
            "color": COLOR_DARK, "align": WD_ALIGN_PARAGRAPH.JUSTIFY
        })

    content_lines.append({"text": f"Window Fee: {_fmt_money(fee)}", "color": COLOR_DARK,
                          "align": WD_ALIGN_PARAGRAPH.JUSTIFY})
    content_lines.append({"text": f"Incidentals Fee: {_fmt_money(incidentals)}", "color": COLOR_DARK,
                          "align": WD_ALIGN_PARAGRAPH.JUSTIFY})
    content_lines.append({"text": f"Total Fees to be received: {_fmt_money(total)}", "color": COLOR_DARK,
                          "bold": True, "align": WD_ALIGN_PARAGRAPH.JUSTIFY})
    content_lines.append({"text": ""})

    # Additional info
    content_lines.append({
        "text": "Additionally, breakfast, lunch and dinner will be provided by the club at your hotel "
                "as per the dates of your assigned games.",
        "color": COLOR_DARK, "align": WD_ALIGN_PARAGRAPH.JUSTIFY
    })
    content_lines.append({"text": ""})
    content_lines.append({
        "text": "Payment for this assignment will be made within 21-days of the window conclusion. ",
        "color": COLOR_DARK, "align": WD_ALIGN_PARAGRAPH.JUSTIFY
    })
    content_lines.append({"text": ""})

    if variant == "F4":
        content_lines.append({
            "text": "If your banking information has recently changed, please be sure to send this "
                    "information to payments.americas@fiba.basketball before the start of the window.",
            "color": COLOR_DARK, "align": WD_ALIGN_PARAGRAPH.JUSTIFY
        })
    else:
        content_lines.append({
            "text": "If your banking information has recently changed, please be sure to send this "
                    "information to payments.americas@fiba.basketball.",
            "color": COLOR_DARK, "align": WD_ALIGN_PARAGRAPH.JUSTIFY
        })

    content_lines.append({"text": ""})
    content_lines.append({
        "text": "If you have any questions, please do not hesitate to contact.",
        "color": COLOR_DARK, "align": WD_ALIGN_PARAGRAPH.JUSTIFY
    })
    content_lines.append({"text": ""})

    # Insert all content paragraphs into the document before ref_element
    # First, use paragraph [4] for the first content line
    first = content_lines[0]
    if "mixed" in first:
        _set_para_mixed_font(paras[4], first["mixed"], font_name)
    else:
        _set_para_text_font(paras[4], first.get("text", ""), first.get("color", COLOR_DARK),
                            font_name, bold=first.get("bold", False),
                            size=first.get("size"), align=first.get("align"))

    # Insert remaining content paragraphs after [4]
    from docx.oxml.ns import qn as _qn
    from copy import deepcopy

    insert_after = paras[4]._element
    for line in content_lines[1:]:
        new_p = doc.element.makeelement(_qn('w:p'), {})
        insert_after.addnext(new_p)
        insert_after = new_p

        # Create a temporary paragraph wrapper
        from docx.text.paragraph import Paragraph
        para = Paragraph(new_p, doc)

        if line.get("align") is not None:
            para.alignment = line["align"]

        text = line.get("text", "")
        if "mixed" in line:
            for t, color, bold in line["mixed"]:
                run = para.add_run(t)
                run.font.name = font_name
                run.font.color.rgb = color
                run.bold = bold
                run.font.size = Pt(10)
        elif text:
            run = para.add_run(text)
            run.font.name = font_name
            run.font.color.rgb = line.get("color", COLOR_DARK)
            run.bold = line.get("bold", False)
            run.font.size = line.get("size", Pt(10))
        # else: empty paragraph, leave as-is

    return doc


# ─── CONFIRMATION (BCLA / LSB) FROM SCRATCH ──────────────────────────────────

def _build_confirmation_from_scratch(data: dict) -> Document:
    doc = Document()
    _apply_base_style(doc)

    nominee = data.get("nominee_name", "")
    comp_name = data.get("competition_name", "")
    role = data.get("role", "VGO")
    role_label = "Video Graphic Operator" if role == "VGO" else "Technical Delegate"
    tk = data.get("template_key", "BCLA")
    game_dates = data.get("game_dates") or []

    title = f"Confirmation – {comp_name}"
    if tk == "LSB":
        title += f" {data.get('competition_year', '')}"

    _add_heading(doc, title)
    _add_empty(doc)
    _add_body(doc, [("Dear ", COLOR_DARK), (nominee, COLOR_RED), (",", COLOR_DARK)])
    _add_empty(doc)
    _add_body_text(doc, f"This letter confirms your assignment as {role_label} for the {comp_name}.")
    _add_empty(doc)

    details = []
    if data.get("location"):
        details.append(f"Location: {data['location']}")
    if data.get("venue"):
        details.append(f"Venue: {data['venue']}")
    if data.get("arrival_date"):
        details.append(f"Arrival Date: {_fmt_date(data['arrival_date'])}")
    if data.get("departure_date"):
        details.append(f"Departure Date: {_fmt_date(data['departure_date'])}")
    for d in details:
        _add_body_text(doc, f"  •  {d}")
    _add_empty(doc)

    for gd in game_dates:
        label = gd.get("label", "")
        date_val = _fmt_date(gd.get("date", ""))
        text = f"{label}: {date_val}" if label else date_val
        _add_centered_red(doc, text)
    _add_empty(doc)

    _add_body_text(doc, f"Below list the details of payment you will receive as {role_label} assigned to the competition listed above:")
    _add_empty(doc)
    _add_fee_line(doc, f"Window Fee: {_fmt_money(data.get('window_fee'))}", bold=False)
    _add_fee_line(doc, f"Incidentals: {_fmt_money(data.get('incidentals'))}", bold=False)
    _add_fee_line(doc, f"Total: {_fmt_money(data.get('total'))}", bold=True)
    _add_empty(doc)

    _add_body_text(doc, "Thank you for your commitment and professionalism.")
    _add_empty(doc)
    _add_empty(doc)

    sig_name, sig_title, sig_org = SIGNATORIES.get(tk, SIGNATORIES["BCLA"])
    _add_body_text(doc, f"{sig_name} {sig_title} {sig_org}")

    return doc


# ─── WCQ FROM SCRATCH (fallback) ─────────────────────────────────────────────

def _build_wcq_from_scratch(data: dict) -> Document:
    doc = Document()
    _apply_base_style(doc)

    nominee = data.get("nominee_name", "")
    comp_name = data.get("competition_name", "")
    role = data.get("role", "VGO")
    role_label = "Video Graphic Operator" if role == "VGO" else "Technical Delegate"
    game_dates = data.get("game_dates") or []
    deadline = _fmt_deadline(data.get("confirmation_deadline", ""))

    _add_heading(doc, f"Nomination for the {comp_name}")
    _add_empty(doc)
    _add_body(doc, [("Dear ", COLOR_DARK), (nominee, COLOR_RED), (",", COLOR_DARK)])
    _add_empty(doc)
    _add_body_text(doc, f"We would like to inform that you have been nominated for the following games of the {comp_name}.")
    _add_empty(doc)

    for gd in game_dates:
        label = gd.get("label", "")
        date_val = _fmt_date(gd.get("date", ""))
        text = f"{label}: {date_val}" if label else date_val
        _add_centered_red(doc, text)

    _add_empty(doc)
    _add_empty(doc)

    confirm_email = CONFIRMATION_EMAIL.get(role, CONFIRMATION_EMAIL["VGO"])
    parts = [
        (f"As per the FIBA Internal Regulations Book 3, please confirm to us your availability to fulfil your assignment as {role_label} by ", COLOR_DARK),
        (deadline, COLOR_RED),
        (f". Confirmation shall be sent to {confirm_email}", COLOR_DARK),
    ]
    _add_body(doc, parts, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    _add_empty(doc)
    _add_body_text(doc, "As soon as we receive your confirmation, we will make arrangements for international flights to the host country and provide you with relevant information in order for you to prepare the game and establish contact with the Game Director of the Host National Federation.")
    _add_empty(doc)
    _add_body_text(doc, f"Below list the details of payment you will receive as {role_label} assigned to the competition listed above:")
    _add_empty(doc)
    _add_empty(doc)
    _add_empty(doc)

    _add_fee_line(doc, f"Per Game Fee: {_fmt_money(data.get('window_fee'))}", bold=False)
    _add_fee_line(doc, f"Incidentals: {_fmt_money(data.get('incidentals'))}", bold=False)
    _add_fee_line(doc, f"Total: {_fmt_money(data.get('total'))}", bold=True)

    _add_empty(doc)
    _add_empty(doc)
    _add_empty(doc)

    _add_body_text(doc, "We wish you the best in your preparation and accomplishment of your assignment.")
    _add_empty(doc)
    _add_empty(doc)

    sig_name, sig_title, sig_org = SIGNATORIES.get(data.get("template_key", "GENERIC"), SIGNATORIES["GENERIC"])
    _add_body_text(doc, f"{sig_name} {sig_title} {sig_org}")

    return doc


# ─── PDF CONVERSION (CloudConvert) ───────────────────────────────────────────

def _convert_to_pdf(docx_path: str) -> tuple[str | None, str | None]:
    """
    Convert .docx to .pdf using CloudConvert API.
    Returns (pdf_path, error_message). If conversion succeeds, error is None.
    """
    api_key = os.environ.get("CLOUDCONVERT_API_KEY", "").strip()
    if not api_key:
        return None, "CLOUDCONVERT_API_KEY not set"

    pdf_path = docx_path.replace(".docx", ".pdf")
    base_url = "https://api.cloudconvert.com/v2"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }

    try:
        # Step 1: Create job with import + convert + export
        job_payload = {
            "tasks": {
                "import-file": {
                    "operation": "import/upload",
                },
                "convert-file": {
                    "operation": "convert",
                    "input": ["import-file"],
                    "output_format": "pdf",
                    "engine": "libreoffice",
                },
                "export-file": {
                    "operation": "export/url",
                    "input": ["convert-file"],
                },
            }
        }

        job_resp = httpx.post(
            f"{base_url}/jobs", json=job_payload, headers=headers, timeout=30.0
        )
        if job_resp.status_code not in (200, 201):
            err = f"Job create error {job_resp.status_code}: {job_resp.text[:300]}"
            print(f"[CLOUDCONVERT] {err}")
            return None, err

        job_data = job_resp.json()["data"]

        # Step 2: Find the upload task and upload the file
        upload_task = None
        for task in job_data["tasks"]:
            if task["name"] == "import-file" and task.get("result", {}).get("form"):
                upload_task = task
                break

        if not upload_task:
            return None, "No upload task found in job response"

        form_data = upload_task["result"]["form"]
        upload_url = form_data["url"]
        form_params = form_data["parameters"]

        with open(docx_path, "rb") as f:
            docx_bytes = f.read()

        filename = Path(docx_path).name

        # Build multipart upload
        files = {"file": (filename, docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        upload_resp = httpx.post(
            upload_url, data=form_params, files=files, timeout=60.0
        )
        if upload_resp.status_code not in (200, 201, 204):
            err = f"Upload error {upload_resp.status_code}: {upload_resp.text[:200]}"
            print(f"[CLOUDCONVERT] {err}")
            return None, err

        # Step 3: Wait for job to complete (poll)
        job_id = job_data["id"]
        import time
        status_data = None
        for _ in range(30):
            time.sleep(1)
            status_resp = httpx.get(
                f"{base_url}/jobs/{job_id}", headers=headers, timeout=15.0
            )
            if status_resp.status_code != 200:
                continue
            status_data = status_resp.json()["data"]
            if status_data["status"] == "finished":
                break
            elif status_data["status"] == "error":
                err = f"Job failed: {status_data}"
                print(f"[CLOUDCONVERT] {err}")
                return None, err

        if not status_data or status_data["status"] != "finished":
            return None, f"Job timed out. Last status: {status_data.get('status') if status_data else 'unknown'}"

        # Step 4: Get export URL and download PDF
        export_task = None
        for task in status_data["tasks"]:
            if task["name"] == "export-file" and task["status"] == "finished":
                export_task = task
                break

        if not export_task or not export_task.get("result", {}).get("files"):
            return None, "No export result found"

        download_url = export_task["result"]["files"][0]["url"]
        pdf_resp = httpx.get(download_url, timeout=30.0)
        if pdf_resp.status_code == 200:
            with open(pdf_path, "wb") as f:
                f.write(pdf_resp.content)
            return pdf_path, None
        else:
            return None, f"Download error {pdf_resp.status_code}"

    except Exception as e:
        import traceback
        traceback.print_exc()
        return None, f"{type(e).__name__}: {e}"


# ─── PARAGRAPH HELPERS ───────────────────────────────────────────────────────

def _set_para_text(para, text, color, bold=False, size=None, align=None):
    _clear_para(para)
    if align is not None:
        para.alignment = align
    run = para.add_run(text)
    run.font.name = FONT_NAME
    run.font.color.rgb = color
    run.bold = bold
    if size:
        run.font.size = size


def _set_para_mixed(para, parts, align=None):
    _clear_para(para)
    if align is not None:
        para.alignment = align
    for text, color, bold in parts:
        run = para.add_run(text)
        run.font.name = FONT_NAME
        run.font.color.rgb = color
        run.bold = bold


def _clear_para(para):
    for r in para._element.findall(qn('w:r')):
        para._element.remove(r)


def _apply_base_style(doc):
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(10)
    style.font.color.rgb = COLOR_DARK


def _add_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = COLOR_DARK


def _add_body_text(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_DARK


def _add_body(doc, parts, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    for text, color in parts:
        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = Pt(10)
        run.font.color.rgb = color


def _add_centered_red(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_RED


def _add_fee_line(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_RED
    run.bold = bold


def _add_empty(doc):
    doc.add_paragraph()


def _clear_para_generic(para):
    """Clear paragraph runs (same as _clear_para but also checks for drawings to preserve)."""
    for r in para._element.findall(qn('w:r')):
        # Skip runs that contain drawing elements
        if r.findall(qn('w:drawing')):
            continue
        para._element.remove(r)


def _set_para_text_font(para, text, color, font_name, bold=False, size=None, align=None):
    _clear_para_generic(para)
    if align is not None:
        para.alignment = align
    run = para.add_run(text)
    run.font.name = font_name
    run.font.color.rgb = color
    run.bold = bold
    if size:
        run.font.size = size


def _set_para_mixed_font(para, parts, font_name, align=None):
    _clear_para_generic(para)
    if align is not None:
        para.alignment = align
    for text, color, bold in parts:
        run = para.add_run(text)
        run.font.name = font_name
        run.font.color.rgb = color
        run.bold = bold


def _fmt_money(val) -> str:
    if val is None:
        return ""
    try:
        v = float(val)
        if v == int(v):
            return f"${int(v)}"
        return f"${v:,.2f}"
    except (ValueError, TypeError):
        return str(val)


# ─── STORAGE ─────────────────────────────────────────────────────────────────

def _upload_to_storage(file_path: str, base_name: str) -> str | None:
    """Upload generated file to Supabase Storage bucket 'nominations'."""
    import traceback
    try:
        from api._lib.database import get_supabase

        client = get_supabase()
        bucket_name = "nominations"
        ext = Path(file_path).suffix
        storage_path = f"{base_name}{ext}"

        with open(file_path, "rb") as f:
            file_bytes = f.read()

        content_type = "application/pdf" if ext == ".pdf" else \
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        try:
            client.storage.from_(bucket_name).upload(
                path=storage_path,
                file=file_bytes,
                file_options={"content-type": content_type, "upsert": "true"},
            )
        except Exception:
            try:
                client.storage.from_(bucket_name).remove([storage_path])
            except Exception:
                pass
            client.storage.from_(bucket_name).upload(
                path=storage_path,
                file=file_bytes,
                file_options={"content-type": content_type},
            )

        public_url = client.storage.from_(bucket_name).get_public_url(storage_path)
        return public_url

    except Exception as e:
        print(f"[STORAGE ERROR] {type(e).__name__}: {e}")
        traceback.print_exc()
        return None
