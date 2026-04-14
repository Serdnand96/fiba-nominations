from fastapi import APIRouter, HTTPException, Query, UploadFile, File, Form
from pydantic import BaseModel
from typing import Optional, List
from datetime import datetime, timedelta
import io
import tempfile
import os

from api._lib.database import supabase

router = APIRouter(prefix="/training", tags=["training"])


# ── Schemas ──────────────────────────────────────────────────────────────────

class SlotCreate(BaseModel):
    competition_id: str
    date: str          # YYYY-MM-DD
    start_time: str    # HH:MM
    end_time: str      # HH:MM
    venue: str
    team_label: str
    sport: str = "Basketball"
    notes: Optional[str] = None


class SlotUpdate(BaseModel):
    date: Optional[str] = None
    start_time: Optional[str] = None
    end_time: Optional[str] = None
    venue: Optional[str] = None
    team_label: Optional[str] = None
    sport: Optional[str] = None
    notes: Optional[str] = None


class BulkSlotCreate(BaseModel):
    competition_id: str
    slots: List[SlotCreate]


class AssignmentCreate(BaseModel):
    training_slot_id: str
    personnel_id: str


# ── Helpers ──────────────────────────────────────────────────────────────────

def _enrich_slots(slots: list) -> list:
    """Add assigned TDs info to each slot."""
    if not slots:
        return slots

    slot_ids = [s["id"] for s in slots]

    # Fetch all assignments for these slots
    all_assignments = []
    for sid in slot_ids:
        asn = supabase.table("training_slot_assignments").select("*").eq("training_slot_id", sid).execute().data
        all_assignments.extend(asn)

    if not all_assignments:
        for s in slots:
            s["assignments"] = []
        return slots

    # Fetch personnel info
    personnel_ids = list({a["personnel_id"] for a in all_assignments})
    personnel_map = {}
    for pid in personnel_ids:
        p = supabase.table("personnel").select("id,name,country,role").eq("id", pid).execute().data
        if p:
            personnel_map[pid] = p[0]

    # Group assignments by slot
    slot_assignments = {}
    for a in all_assignments:
        sid = a["training_slot_id"]
        if sid not in slot_assignments:
            slot_assignments[sid] = []
        a["personnel"] = personnel_map.get(a["personnel_id"])
        slot_assignments[sid].append(a)

    for s in slots:
        s["assignments"] = slot_assignments.get(s["id"], [])

    return slots


def _check_td_conflicts(personnel_id: str, target_date: str, target_start: str, target_end: str, exclude_slot_id: str = None):
    """Check if a TD has overlapping training slots on the same date."""
    # Get all slots assigned to this TD on the same date
    all_assignments = supabase.table("training_slot_assignments").select("training_slot_id").eq("personnel_id", personnel_id).execute().data

    if not all_assignments:
        return {"has_conflict": False, "conflict_detail": None}

    slot_ids = [a["training_slot_id"] for a in all_assignments]

    conflicts = []
    for sid in slot_ids:
        if sid == exclude_slot_id:
            continue
        slot = supabase.table("training_slots").select("*").eq("id", sid).execute().data
        if not slot:
            continue
        slot = slot[0]
        if slot["date"] != target_date:
            continue

        # Check time overlap
        if _times_overlap(slot["start_time"], slot["end_time"], target_start, target_end):
            conflicts.append(f"{slot['team_label']} ({slot['start_time'][:5]}-{slot['end_time'][:5]}) @ {slot['venue']}")

    if conflicts:
        return {
            "has_conflict": True,
            "conflict_detail": "; ".join(conflicts),
        }

    return {"has_conflict": False, "conflict_detail": None}


def _times_overlap(start_a: str, end_a: str, start_b: str, end_b: str) -> bool:
    """Check if two time ranges overlap on the same day."""
    def to_minutes(t: str) -> int:
        parts = t.split(":")
        return int(parts[0]) * 60 + int(parts[1])

    a_start = to_minutes(start_a)
    a_end = to_minutes(end_a)
    b_start = to_minutes(start_b)
    b_end = to_minutes(end_b)

    return a_start < b_end and b_start < a_end


def _format_time(t: str) -> str:
    """Normalize time string to HH:MM."""
    if not t:
        return t
    parts = t.split(":")
    return f"{parts[0].zfill(2)}:{parts[1].zfill(2)}"


# ── Slots CRUD ───────────────────────────────────────────────────────────────

@router.get("/slots")
def list_slots(competition_id: str = Query(...)):
    """List all training slots for a competition, with assigned TDs."""
    slots = (
        supabase.table("training_slots")
        .select("*")
        .eq("competition_id", competition_id)
        .order("date")
        .execute()
        .data
    )
    # Secondary sort by start_time (PostgREST only supports one order)
    slots.sort(key=lambda s: (s["date"], s["start_time"]))
    return _enrich_slots(slots)


@router.get("/slots/by-team")
def list_slots_by_team(competition_id: str = Query(...), team_label: str = Query(...)):
    slots = (
        supabase.table("training_slots")
        .select("*")
        .eq("competition_id", competition_id)
        .eq("team_label", team_label)
        .order("date")
        .execute()
        .data
    )
    slots.sort(key=lambda s: (s["date"], s["start_time"]))
    return _enrich_slots(slots)


@router.get("/slots/by-date")
def list_slots_by_date(competition_id: str = Query(...), date: str = Query(...)):
    slots = (
        supabase.table("training_slots")
        .select("*")
        .eq("competition_id", competition_id)
        .eq("date", date)
        .order("date")
        .execute()
        .data
    )
    slots.sort(key=lambda s: s["start_time"])
    return _enrich_slots(slots)


@router.get("/slots/by-personnel")
def list_slots_by_personnel(personnel_id: str = Query(...), competition_id: str = Query(...)):
    """Get all training slots assigned to a specific TD in a competition."""
    # Get all assignments for this person
    assignments = (
        supabase.table("training_slot_assignments")
        .select("training_slot_id")
        .eq("personnel_id", personnel_id)
        .execute()
        .data
    )

    if not assignments:
        return []

    slot_ids = [a["training_slot_id"] for a in assignments]

    # Fetch the slots that belong to this competition
    slots = []
    for sid in slot_ids:
        s = (
            supabase.table("training_slots")
            .select("*")
            .eq("id", sid)
            .eq("competition_id", competition_id)
            .execute()
            .data
        )
        if s:
            slots.append(s[0])

    slots.sort(key=lambda s: (s["date"], s["start_time"]))
    return _enrich_slots(slots)


@router.post("/slots")
def create_slot(data: SlotCreate):
    record = data.model_dump()
    record["start_time"] = _format_time(record["start_time"])
    record["end_time"] = _format_time(record["end_time"])
    result = supabase.table("training_slots").insert(record).execute()
    return result.data[0]


@router.put("/slots/{slot_id}")
def update_slot(slot_id: str, data: SlotUpdate):
    updates = {k: v for k, v in data.model_dump().items() if v is not None}
    if not updates:
        raise HTTPException(400, "No fields to update")
    if "start_time" in updates:
        updates["start_time"] = _format_time(updates["start_time"])
    if "end_time" in updates:
        updates["end_time"] = _format_time(updates["end_time"])
    updates["updated_at"] = datetime.utcnow().isoformat()
    r = supabase.table("training_slots").update(updates).eq("id", slot_id).execute()
    if not r.data:
        raise HTTPException(404, "Slot not found")
    return r.data[0]


@router.delete("/slots/{slot_id}")
def delete_slot(slot_id: str):
    supabase.table("training_slots").delete().eq("id", slot_id).execute()
    return {"ok": True}


@router.post("/slots/bulk")
def bulk_create_slots(data: BulkSlotCreate):
    """Bulk create/upsert slots. Dedup by (competition_id, date, start_time, team_label)."""
    created = 0
    skipped = 0
    errors = []

    for slot in data.slots:
        try:
            record = slot.model_dump()
            record["competition_id"] = data.competition_id
            record["start_time"] = _format_time(record["start_time"])
            record["end_time"] = _format_time(record["end_time"])

            # Check for existing slot
            existing = (
                supabase.table("training_slots")
                .select("id")
                .eq("competition_id", data.competition_id)
                .eq("date", record["date"])
                .eq("start_time", record["start_time"])
                .eq("team_label", record["team_label"])
                .execute()
                .data
            )

            if existing:
                # Update existing
                supabase.table("training_slots").update({
                    "end_time": record["end_time"],
                    "venue": record["venue"],
                    "sport": record.get("sport", "Basketball"),
                    "notes": record.get("notes"),
                    "updated_at": datetime.utcnow().isoformat(),
                }).eq("id", existing[0]["id"]).execute()
                skipped += 1
            else:
                supabase.table("training_slots").insert(record).execute()
                created += 1
        except Exception as e:
            errors.append({"team_label": slot.team_label, "date": slot.date, "error": str(e)})

    return {"imported": created, "updated": skipped, "errors": errors}


# ── Assignments ──────────────────────────────────────────────────────────────

@router.post("/assignments")
def create_assignment(data: AssignmentCreate):
    """Assign a TD to a training slot. Returns conflict warning if any."""
    # Get the slot details for conflict check
    slot = supabase.table("training_slots").select("*").eq("id", data.training_slot_id).execute().data
    if not slot:
        raise HTTPException(404, "Training slot not found")
    slot = slot[0]

    # Check conflicts
    conflict = _check_td_conflicts(
        data.personnel_id,
        slot["date"],
        slot["start_time"],
        slot["end_time"],
        exclude_slot_id=None,
    )

    # Always create the assignment (conflict is a warning, not a blocker)
    try:
        result = supabase.table("training_slot_assignments").insert(data.model_dump()).execute()
        assignment = result.data[0]
    except Exception as e:
        if "unique" in str(e).lower() or "duplicate" in str(e).lower():
            raise HTTPException(409, "TD already assigned to this slot")
        raise HTTPException(500, str(e))

    # Enrich with personnel info
    person = supabase.table("personnel").select("id,name,country,role").eq("id", data.personnel_id).execute().data
    if person:
        assignment["personnel"] = person[0]

    return {
        **assignment,
        "conflict_warning": conflict["conflict_detail"] if conflict["has_conflict"] else None,
    }


@router.delete("/assignments/{assignment_id}")
def delete_assignment(assignment_id: str):
    supabase.table("training_slot_assignments").delete().eq("id", assignment_id).execute()
    return {"ok": True}


# ── Conflict check endpoint ─────────────────────────────────────────────────

@router.get("/conflicts")
def check_conflicts(personnel_id: str = Query(...), slot_id: str = Query(...)):
    """Check if a TD would have conflicts if assigned to a slot."""
    slot = supabase.table("training_slots").select("*").eq("id", slot_id).execute().data
    if not slot:
        raise HTTPException(404, "Slot not found")
    slot = slot[0]

    return _check_td_conflicts(personnel_id, slot["date"], slot["start_time"], slot["end_time"])


# ── Excel import ─────────────────────────────────────────────────────────────

@router.post("/import/excel")
async def import_excel(
    file: UploadFile = File(...),
    competition_id: str = Form(...),
    sport: str = Form("Basketball"),
):
    """Import training slots from FIBA multi-sport schedule Excel."""
    try:
        import openpyxl
    except ImportError:
        # Fallback: try pandas
        pass

    content = await file.read()

    try:
        slots = _parse_fiba_schedule(content, competition_id, sport)
    except Exception as e:
        raise HTTPException(400, f"Error parsing Excel file: {str(e)}")

    if not slots:
        return {"imported": 0, "skipped": 0, "errors": [], "preview": []}

    # Bulk insert with dedup
    created = 0
    skipped = 0
    errors = []

    for slot in slots:
        try:
            existing = (
                supabase.table("training_slots")
                .select("id")
                .eq("competition_id", competition_id)
                .eq("date", slot["date"])
                .eq("start_time", slot["start_time"])
                .eq("team_label", slot["team_label"])
                .execute()
                .data
            )

            if existing:
                supabase.table("training_slots").update({
                    "end_time": slot["end_time"],
                    "venue": slot["venue"],
                    "sport": slot["sport"],
                    "updated_at": datetime.utcnow().isoformat(),
                }).eq("id", existing[0]["id"]).execute()
                skipped += 1
            else:
                supabase.table("training_slots").insert(slot).execute()
                created += 1
        except Exception as e:
            errors.append({"team_label": slot.get("team_label"), "error": str(e)})

    return {"imported": created, "skipped": skipped, "errors": errors}


def _parse_fiba_schedule(content: bytes, competition_id: str, sport: str) -> list:
    """Parse FIBA multi-sport schedule Excel format."""
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
    ws = wb.active

    slots = []
    current_date = None

    for row in ws.iter_rows(min_row=1, values_only=False):
        cells = [c.value for c in row]

        # Detect date header row: "Fecha" in first columns
        for i, val in enumerate(cells[:3]):
            if val and str(val).strip().upper() == "FECHA":
                # Date is in the next cell(s)
                for j in range(i + 1, min(len(cells), i + 5)):
                    if cells[j]:
                        date_val = cells[j]
                        if isinstance(date_val, datetime):
                            current_date = date_val.strftime("%Y-%m-%d")
                        elif isinstance(date_val, str):
                            current_date = date_val.strip()
                        break
                break

        if not current_date:
            continue

        # Get time from column C (index 2)
        time_val = cells[2] if len(cells) > 2 else None
        if not time_val:
            continue

        # Skip "PARTIDOS" rows
        if isinstance(time_val, str) and "PARTIDOS" in str(time_val).upper():
            continue

        # Skip non-time values like "Comienza", "DIA", headers
        time_str = None
        if isinstance(time_val, datetime):
            time_str = time_val.strftime("%H:%M")
        elif isinstance(time_val, str):
            # Try to parse HH:MM:SS or HH:MM
            stripped = time_val.strip()
            parts = stripped.split(":")
            if len(parts) >= 2:
                try:
                    h = int(parts[0])
                    m = int(parts[1])
                    if 0 <= h <= 23 and 0 <= m <= 59:
                        time_str = f"{h:02d}:{m:02d}"
                except (ValueError, IndexError):
                    pass

        if not time_str:
            continue

        # Calculate end time (start + 90 minutes)
        h, m = map(int, time_str.split(":"))
        end_minutes = h * 60 + m + 90
        end_h = end_minutes // 60
        end_m = end_minutes % 60
        end_time_str = f"{end_h:02d}:{end_m:02d}"

        # Check columns for team labels
        # Column F (index 5) or H (index 7) = Estadio
        # Column H (index 7) or I (index 8) = Cancha de Entrenamiento
        # Try multiple column layouts since Excel files vary
        estadio_cols = [5, 7]
        cancha_cols = [8]

        for col_idx in estadio_cols:
            if col_idx < len(cells) and cells[col_idx]:
                label = str(cells[col_idx]).strip()
                if label and label not in ("Estadio", "Cancha de Entrenamiento", "PARTIDOS", "Comienza"):
                    slots.append({
                        "competition_id": competition_id,
                        "date": current_date,
                        "start_time": time_str,
                        "end_time": end_time_str,
                        "venue": "Estadio",
                        "team_label": label,
                        "sport": sport,
                    })

        for col_idx in cancha_cols:
            if col_idx < len(cells) and cells[col_idx]:
                label = str(cells[col_idx]).strip()
                if label and label not in ("Estadio", "Cancha de Entrenamiento", "PARTIDOS", "Comienza"):
                    slots.append({
                        "competition_id": competition_id,
                        "date": current_date,
                        "start_time": time_str,
                        "end_time": end_time_str,
                        "venue": "Cancha de Entrenamiento",
                        "team_label": label,
                        "sport": sport,
                    })

    return slots


# ── Excel preview (parse without importing) ──────────────────────────────────

@router.post("/import/preview")
async def preview_excel(
    file: UploadFile = File(...),
    competition_id: str = Form(...),
    sport: str = Form("Basketball"),
):
    """Preview parsed slots from Excel without importing."""
    content = await file.read()
    try:
        slots = _parse_fiba_schedule(content, competition_id, sport)
    except Exception as e:
        raise HTTPException(400, f"Error parsing Excel file: {str(e)}")

    return {"total": len(slots), "preview": slots[:10]}


# ── PDF export ───────────────────────────────────────────────────────────────

@router.get("/export/pdf/competition/{competition_id}")
def export_pdf_competition(competition_id: str):
    """Full competition training schedule PDF."""
    # Get competition info
    comp = supabase.table("competitions").select("*").eq("id", competition_id).execute().data
    if not comp:
        raise HTTPException(404, "Competition not found")
    comp = comp[0]

    slots = (
        supabase.table("training_slots")
        .select("*")
        .eq("competition_id", competition_id)
        .order("date")
        .execute()
        .data
    )
    slots.sort(key=lambda s: (s["date"], s["start_time"]))
    slots = _enrich_slots(slots)

    title = f"Training Schedule - {comp['name']}"
    subtitle = f"All dates"
    return _generate_schedule_pdf(title, subtitle, slots)


@router.get("/export/pdf/daily")
def export_pdf_daily(competition_id: str = Query(...), date: str = Query(...)):
    comp = supabase.table("competitions").select("*").eq("id", competition_id).execute().data
    if not comp:
        raise HTTPException(404, "Competition not found")
    comp = comp[0]

    slots = (
        supabase.table("training_slots")
        .select("*")
        .eq("competition_id", competition_id)
        .eq("date", date)
        .order("date")
        .execute()
        .data
    )
    slots.sort(key=lambda s: s["start_time"])
    slots = _enrich_slots(slots)

    title = f"Training Schedule - {comp['name']}"
    subtitle = f"Date: {date}"
    return _generate_schedule_pdf(title, subtitle, slots)


@router.get("/export/pdf/team")
def export_pdf_team(competition_id: str = Query(...), team_label: str = Query(...)):
    comp = supabase.table("competitions").select("*").eq("id", competition_id).execute().data
    if not comp:
        raise HTTPException(404, "Competition not found")
    comp = comp[0]

    slots = (
        supabase.table("training_slots")
        .select("*")
        .eq("competition_id", competition_id)
        .eq("team_label", team_label)
        .order("date")
        .execute()
        .data
    )
    slots.sort(key=lambda s: (s["date"], s["start_time"]))
    slots = _enrich_slots(slots)

    title = f"Training Schedule - {comp['name']}"
    subtitle = f"Team: {team_label}"
    return _generate_schedule_pdf(title, subtitle, slots)


def _generate_schedule_pdf(title: str, subtitle: str, slots: list):
    """Generate a training schedule PDF using python-docx, return as download."""
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from fastapi.responses import FileResponse

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 51, 102)

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # Generation date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(150, 150, 150)

    doc.add_paragraph()  # spacer

    if not slots:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("No training slots found.")
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(150, 150, 150)
    else:
        # Table
        table = doc.add_table(rows=1, cols=6)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = ["Date", "Start", "End", "Venue", "Team", "Assigned TDs"]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)

        for slot in slots:
            row = table.add_row()
            tds = ", ".join(
                a["personnel"]["name"]
                for a in slot.get("assignments", [])
                if a.get("personnel")
            ) or "-"

            values = [
                slot["date"],
                slot["start_time"][:5],
                slot["end_time"][:5],
                slot["venue"],
                slot["team_label"],
                tds,
            ]
            for i, val in enumerate(values):
                cell = row.cells[i]
                cell.text = str(val)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

    # Save to temp file
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    tmp.close()

    # Try CloudConvert PDF conversion
    pdf_path = _convert_to_pdf(tmp.name)

    if pdf_path:
        os.unlink(tmp.name)
        return FileResponse(
            pdf_path,
            media_type="application/pdf",
            filename="training_schedule.pdf",
        )
    else:
        return FileResponse(
            tmp.name,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename="training_schedule.docx",
        )


def _convert_to_pdf(docx_path: str) -> str | None:
    """Convert docx to PDF using CloudConvert. Returns PDF path or None."""
    import httpx
    import time

    api_key = os.environ.get("CLOUDCONVERT_API_KEY", "").strip()
    if not api_key:
        return None

    try:
        base_url = "https://api.cloudconvert.com/v2"
        hdrs = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}

        job_payload = {"tasks": {
            "import-file": {"operation": "import/upload"},
            "convert-file": {"operation": "convert", "input": ["import-file"], "output_format": "pdf", "engine": "libreoffice"},
            "export-file": {"operation": "export/url", "input": ["convert-file"]},
        }}

        job_resp = httpx.post(f"{base_url}/jobs", json=job_payload, headers=hdrs, timeout=30.0)
        if job_resp.status_code not in (200, 201):
            return None

        job_data = job_resp.json()["data"]
        upload_task = next(
            (t for t in job_data["tasks"] if t["name"] == "import-file" and t.get("result", {}).get("form")),
            None,
        )
        if not upload_task:
            return None

        form_d = upload_task["result"]["form"]
        with open(docx_path, "rb") as f:
            docx_bytes = f.read()

        upload_resp = httpx.post(
            form_d["url"],
            data=form_d["parameters"],
            files={"file": ("schedule.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            timeout=60.0,
        )
        if upload_resp.status_code not in (200, 201, 204):
            return None

        job_id = job_data["id"]
        status_data = None
        for _ in range(30):
            time.sleep(1)
            sr = httpx.get(f"{base_url}/jobs/{job_id}", headers=hdrs, timeout=15.0)
            if sr.status_code == 200:
                status_data = sr.json()["data"]
                if status_data["status"] in ("finished", "error"):
                    break

        if not status_data or status_data["status"] != "finished":
            return None

        export_task = next(
            (t for t in status_data["tasks"] if t["name"] == "export-file" and t["status"] == "finished"),
            None,
        )
        if not export_task or not export_task.get("result", {}).get("files"):
            return None

        dl_url = export_task["result"]["files"][0]["url"]
        pdf_resp = httpx.get(dl_url, timeout=30.0)

        pdf_path = docx_path.replace(".docx", ".pdf")
        with open(pdf_path, "wb") as f:
            f.write(pdf_resp.content)

        return pdf_path
    except Exception:
        return None
